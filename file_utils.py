# file_utils.py
import os
import zipfile
import csv
import docx
import pdfplumber
from fpdf import FPDF

# ----------------- File Compression -----------------
def compress_file(input_file, output_file):
    """Compress a single file into a ZIP archive"""
    try:
        with zipfile.ZipFile(output_file, 'w') as zipf:
            zipf.write(input_file, os.path.basename(input_file))
        return f"✅ Successfully compressed {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in compression: {e}"

# ----------------- DOCX <-> PDF <-> CSV <-> TXT -----------------

# DOCX to PDF
def convert_docx_to_pdf(input_file, output_file):
    try:
        doc = docx.Document(input_file)
        pdf = FPDF()
        pdf.add_page()
        for para in doc.paragraphs:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, para.text)
        pdf.output(output_file)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# DOCX to CSV
def convert_docx_to_csv(input_file, output_file):
    try:
        doc = docx.Document(input_file)
        rows = [[para.text] for para in doc.paragraphs if para.text.strip()]
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerows(rows)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# CSV to DOCX
def convert_csv_to_docx(input_file, output_file):
    try:
        doc = docx.Document()
        with open(input_file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                doc.add_paragraph(', '.join(row))
        doc.save(output_file)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# CSV to PDF
def convert_csv_to_pdf(input_file, output_file):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        with open(input_file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                pdf.multi_cell(0, 10, ', '.join(row))
        pdf.output(output_file)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# PDF to TXT
def convert_pdf_to_txt(input_file, output_file):
    try:
        text = ""
        with pdfplumber.open(input_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# PDF to CSV
def convert_pdf_to_csv(input_file, output_file):
    try:
        rows = []
        with pdfplumber.open(input_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        rows.append([line])
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerows(rows)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# TXT to DOCX
def convert_txt_to_docx(input_file, output_file):
    try:
        doc = docx.Document()
        with open(input_file, 'r', encoding='utf-8') as f:
            for line in f:
                doc.add_paragraph(line.strip())
        doc.save(output_file)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# TXT to PDF
def convert_txt_to_pdf(input_file, output_file):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        with open(input_file, 'r', encoding='utf-8') as f:
            for line in f:
                pdf.multi_cell(0, 10, line.strip())
        pdf.output(output_file)
        return f"✅ Successfully converted {input_file} to {output_file}"
    except Exception as e:
        return f"❌ Error in conversion: {e}"

# Helper: detect file format
def get_file_format(file_path):
    _, ext = os.path.splitext(file_path)
    return ext.lower().replace(".", "")
